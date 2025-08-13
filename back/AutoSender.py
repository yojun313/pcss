from back.pcss_temp import PCSSEARCH
import pandas as pd
import os
import gspread
from google.oauth2.service_account import Credentials
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import smtplib
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from datetime import datetime
import json
import re

class AutoSender:
    
    def __init__(self):

        self.test = True
        self.target_year = 2024
        
        self.period_day = 7
        self.period_second = self.period_day * 24 * 60 * 60
        self.period_second = 30
        
        self.CrawlOption = 1
        
        self.spreadsheet_url = "https://docs.google.com/spreadsheets/d/1SsGBT17nzA9ItG8QG73lyHGeIab2C6x616zYwEATQHc/edit?gid=0#gid=0"
        self.user_history_path = os.path.join(os.path.dirname(__file__), 'user_history')
        self.conf_df = pd.read_csv(os.path.join(os.path.dirname(__file__), 'data', 'conf.csv'))
        
        
    def get_spreadsheet_data(self, sheet="Sheet1"):
        
        test_data = [
            {
                'Email': 'moonyojun@naver.com',
                'Networks': "TRUE",
                'Security & Cryptography': "FALSE",
                'Computer Vision & Graphics': "FALSE",
                'Data Mining & Information Retrieval': "FALSE",
                'Machine Learning': "FALSE",
                'Theory': "FALSE",
                'Human-Computer Interaction': "FALSE",
                'Linguistics & Speech': "FALSE",
                'Operating Systems': "FALSE",
                'Computer Architecture': "FALSE",
                'Databases': "FALSE",
                'Programming Languages\r': "FALSE",  # '\r'ì´ í¬í•¨ëœ ê²½ìš° ì£¼ì˜
                'Software Engineering': "FALSE",
                'Embedded & Real-Time Systems': "FALSE",
                'High-Performance Computing': "FALSE",
                'Mobile Computing': "FALSE",
                'Robotics': "FALSE",
                'Artificial Intelligence': "FALSE",
                'Computational Biology': "FALSE",
                'Computational Economics': "FALSE",
                'Design Automation': "FALSE",
                'Visualization': "FALSE"
            },
        ]

        # Google Spreadsheet API ì¸ì¦
        json_file_path = os.path.join(os.path.dirname(__file__), 'data', "lock.json")
        scopes = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]

        credentials = Credentials.from_service_account_file(json_file_path, scopes=scopes)
        gc = gspread.authorize(credentials)

        # Google Spreadsheet ì—´ê¸°
        spreadsheet_url = self.spreadsheet_url
        doc = gc.open_by_url(spreadsheet_url)

        # íŠ¹ì • ì›Œí¬ì‹œíŠ¸ ì„ íƒ
        worksheet = doc.worksheet(sheet)

        # ë°ì´í„°ë¥¼ ê°€ì ¸ì™€ DataFrameìœ¼ë¡œ ë³€í™˜
        data = worksheet.get_all_values()
        df = pd.DataFrame(data[1:], columns=data[0])  # ì²« ë²ˆì§¸ í–‰ì„ ì»¬ëŸ¼ìœ¼ë¡œ ì„¤ì •
        
        dict_list = df.to_dict(orient='records')
        if self.test == True:
            dict_list = test_data
        new_dict_list = []
        
        for dict in dict_list:
            new_dict = {}
            true_list = []
            for key, value in dict.items():
                if value == 'TRUE':
                    new_dict[key] = True
                    true_list.append(key)
                elif value == 'FALSE':
                    new_dict[key] = False
                else:
                    new_dict[key] = value
            filtered_df = self.conf_df[self.conf_df['kind'].isin(true_list)]
            conf_list = filtered_df['conference'].tolist()
            
            new_dict['conf_list'] = conf_list
            new_dict['kind_list'] = true_list
            new_dict_list.append(new_dict) 
        
        return new_dict_list
    
    def send_email(self, receiver, title, text, file_path=None):
        mail_json = os.path.join(os.path.dirname(__file__), 'data', "mail_lock.json")
        
        # JSON íŒŒì¼ì—ì„œ ë©”ì¼ ê³„ì • ì •ë³´ ì½ê¸°
        with open(mail_json, 'r', encoding='utf-8') as jsonfile:
            mail_data = json.load(jsonfile)

        sender = mail_data['sender']
        MailPassword = mail_data['password']

        # ì´ë©”ì¼ ë©”ì‹œì§€ ìƒì„±
        msg = MIMEMultipart()
        msg['Subject'] = title
        msg['From'] = sender
        msg['To'] = receiver

        # ì´ë©”ì¼ ë³¸ë¬¸ ì¶”ê°€
        msg.attach(MIMEText(text, 'plain'))

        # íŒŒì¼ ì²¨ë¶€ (file_pathê°€ ì œê³µëœ ê²½ìš°)
        if file_path and os.path.exists(file_path):
            with open(file_path, 'rb') as f:
                file_attachment = MIMEApplication(f.read(), Name=os.path.basename(file_path))
                file_attachment['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
                msg.attach(file_attachment)

        smtp_server = "smtp.gmail.com"
        smtp_port = 587

        # SMTP ì—°ê²° ë° ë©”ì¼ ë³´ë‚´ê¸°
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender, MailPassword)
            server.sendmail(sender, receiver, msg.as_string())
    
    def main(self):
        while True:
            self.auto_send()
            break
    
    def auto_send(self):
        UserDataList = self.get_spreadsheet_data()
        for UserData in UserDataList:
            conf_list = UserData['conf_list']

            pcssearch_obj = PCSSEARCH(self.CrawlOption, False, self.target_year, self.target_year)
            result_json_path = pcssearch_obj.main(conf_list)
            if os.path.exists(result_json_path):
                os.remove(result_json_path)

            FinalData = pcssearch_obj.FinalData
            
            print(f"{UserData['Email']} í¬ë¡¤ë§ ì™„ë£Œ")
            
            historyData = self.manage_history(self.user_history_path, UserData['Email'], option="GET")
            if historyData != {}:
                titles = [entry["title"] for entry in historyData.values() if "title" in entry]
                
                FilteredData = [entry for entry in FinalData.values() if "title" in entry and entry["title"] not in titles]
                FilteredData = {str(i): entry for i, entry in enumerate(FilteredData)}
            else:
                FilteredData = FinalData
            
            mergedData = list(FinalData.values()) + list(FilteredData.values())
            reindexed_data = {str(i): entry for i, entry in enumerate(mergedData)}                             # history ê°±ì‹ 
            
            self.manage_history(self.user_history_path, UserData['Email'], option='SAVE', data=reindexed_data) # ê¸°ë¡ ì €ì¥
            docx_file = self.make_report(self.user_history_path, FilteredData, UserData)

            self.send_email(
                receiver = UserData['Email'],
                title = f"{datetime.now().strftime("%Y-%m-%d")} PCSS Subscribe Summary",
                text = f'Subscribe => {', '.join(UserData['kind_list'])}',
                file_path = docx_file
            )
            os.remove(docx_file)
            print(f"{UserData['Email']} ë°ì´í„° ì €ì¥ ë° ì „ì†¡ ì™„ë£Œ")


    def manage_history(self, folder_path, email, option, data = None):
        # íŒŒì¼ëª…ì— ì•ˆì „í•œ ë¬¸ìë§Œ ì‚¬ìš© (ì´ë©”ì¼ì—ì„œ íŒŒì¼ëª… ë¶ˆê°€ëŠ¥í•œ ë¬¸ì ì œê±°)
        safe_email = re.sub(r'[^\w\.-]', '_', email)
        json_filename = f"{safe_email}.json"
        json_filepath = os.path.join(folder_path, json_filename)

        if option == 'GET':
            # íŒŒì¼ì´ ì¡´ì¬í•˜ëŠ” ê²½ìš°: JSON ë¶ˆëŸ¬ì˜¤ê¸°
            if os.path.exists(json_filepath):
                with open(json_filepath, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
                return json_data
            
            # íŒŒì¼ì´ ì—†ëŠ” ê²½ìš°: ìƒˆ JSON íŒŒì¼ ìƒì„±
            else:
                return {}
        elif option == "SAVE":
            with open(json_filepath, "w", encoding="utf-8") as file:
                json.dump(data, file, indent=4, ensure_ascii=False)

    def make_report(self, folder_path, data, UserData):
        df = self.conf_df
        conference_to_kind = dict(zip(df["conference"], df["kind"]))

        current_date = datetime.now().strftime("%Y%m%d")
        safe_email = re.sub(r'[^\w\.-]', '_', UserData['Email'])
        docx_filename = f"{current_date}_summary.docx"

        for entry in data.values():
            entry["kind"] = conference_to_kind.get(entry["conference"], "Unknown")

        # Kindë³„ë¡œ ë°ì´í„° ì •ë¦¬
        grouped_by_kind = {}
        for entry in data.values():
            kind = entry["kind"]
            if kind not in grouped_by_kind:
                grouped_by_kind[kind] = {}
            conf = entry["conference"]
            if conf not in grouped_by_kind[kind]:
                grouped_by_kind[kind][conf] = []
            grouped_by_kind[kind][conf].append(entry)


        doc = self.make_docx(grouped_by_kind, UserData)
        # Word ë¬¸ì„œ ì €ì¥
        doc_path = os.path.join(folder_path, docx_filename)
        doc.save(doc_path)
        return doc_path

    def add_hyperlink(self, paragraph, text, url):
        """
        ì£¼ì–´ì§„ ë¬¸ë‹¨(paragraph)ì— í•˜ì´í¼ë§í¬ë¥¼ ì¶”ê°€í•˜ëŠ” í•¨ìˆ˜
        """
        part = paragraph._parent.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                              is_external=True)

        # âœ… w ë„¤ì„ìŠ¤í˜ì´ìŠ¤ë¥¼ ëª…í™•í•˜ê²Œ ì¶”ê°€
        hyperlink = parse_xml(
            f'<w:hyperlink r:id="{r_id}" {nsdecls("w", "r")}>'  # âœ… ë„¤ì„ìŠ¤í˜ì´ìŠ¤ ì¶”ê°€
            f'<w:r><w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr>'
            f'<w:t>{text}</w:t></w:r></w:hyperlink>'
        )

        paragraph._element.append(hyperlink)

    def make_docx(self, grouped_by_kind, UserData):
        doc = Document()
        font_name = "ë§‘ì€ ê³ ë”•"

        styles = doc.styles
        for style_name in ["Normal", "Heading1", "Heading2", "Heading3", "Table Grid"]:
            if style_name in styles:
                style = styles[style_name]
                # í•œê¸€ì´ ì˜¬ë°”ë¥´ê²Œ í‘œì‹œë˜ë„ë¡ ì„¤ì •
                style.font.name = font_name
                style.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

                if style_name == "Normal":
                    style.font.size = Pt(11)
                    style.font.color.rgb = RGBColor(0, 0, 0)
                    style.paragraph_format.space_after = Pt(6)
                elif style_name == "Heading1":
                    style.font.size = Pt(22)
                    style.font.bold = True
                    style.font.color.rgb = RGBColor(0xC8, 0x01, 0x50)  # POSTECH Red
                    style.paragraph_format.space_before = Pt(12)
                    style.paragraph_format.space_after = Pt(6)
                elif style_name == "Heading2":
                    style.font.size = Pt(18)
                    style.font.bold = True
                    style.font.color.rgb = RGBColor(0x00, 0x4F, 0x99)  # ì„¸ë ¨ëœ ë¸”ë£¨ í†¤
                    style.paragraph_format.space_before = Pt(10)
                    style.paragraph_format.space_after = Pt(4)
                elif style_name == "Heading3":
                    style.font.size = Pt(16)
                    style.font.bold = True
                    style.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # ì„¸ë ¨ëœ ê·¸ë¦° í†¤
                    style.paragraph_format.space_before = Pt(8)
                    style.paragraph_format.space_after = Pt(4)
                elif style_name == "Table Grid":
                    style.font.size = Pt(10)
                    style.font.color.rgb = RGBColor(0, 0, 0)
                    style.paragraph_format.space_after = Pt(4)

        doc.add_heading(f"{datetime.now().strftime('%Y-%m-%d')} Conference Summary", level=1)

        doc.add_paragraph(f"- ìˆ˜ì‹ ì¸: {UserData['Email']}")
        doc.add_paragraph(f"- êµ¬ë… ëŒ€ìƒ: {', '.join(UserData['kind_list'])}")
        doc.add_paragraph(f"- í¬ë¡¤ë§ ëŒ€ìƒ ì—°ë„: {self.target_year}")

        # Kind -> Conference ìˆœì„œë¡œ ë¬¸ì„œ ì‘ì„±
        for kind, conferences in grouped_by_kind.items():
            doc.add_heading(f"ğŸ“Œ {kind}", level=1)

            for conf, entries in conferences.items():
                doc.add_heading(f"í•™íšŒ: {conf}", level=2)
                doc.add_heading(f"ì¶œì²˜: {entries[0]['source']}", level=3)

                # í…Œì´ë¸” ìƒì„± (ì—´: Target Author, Title, Authors)
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"

                # í…Œì´ë¸” í—¤ë” ì„¤ì •
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "1ì €ì"
                hdr_cells[1].text = "ì œëª©"
                hdr_cells[2].text = "ì €ì ëª©ë¡"

                # í—¤ë” í°íŠ¸ ì„œì‹ ê°œì„  (ì˜ˆì‹œ: í°ìƒ‰ ë³¼ë“œ ê¸€ì”¨)
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    # (í…Œì´ë¸” ì…€ ë°°ê²½ìƒ‰ ì ìš©ì€ XML ì¡°ì‘ì´ í•„ìš”í•˜ë¯€ë¡œ ìƒëµ)

                for entry in entries:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ", ".join(entry["target_author"]) if entry["target_author"] else "N/A"
                    row_cells[1].text = entry["title"]

                    # ì €ì ì´ë¦„ê³¼ URLì„ ë§¤ì¹­í•˜ì—¬ ì¶”ê°€ (í•˜ì´í¼ë§í¬ í¬í•¨)
                    p = row_cells[2].paragraphs[0]
                    for name, url in zip(entry["author_name"], entry["author_url"]):
                        self.add_hyperlink(p, name, url)
                        p.add_run("\n")

                # í…Œì´ë¸”ê³¼ ë‹¤ìŒ ë‚´ìš© ì‚¬ì´ ê°„ê²© ì¶”ê°€
                doc.add_paragraph("\n")

        return doc


if __name__ == '__main__':
    autoSender_obj = AutoSender()
    autoSender_obj.main()